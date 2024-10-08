import os
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import requests
import logging
from supabase import create_client, Client
from oauthlib.oauth2 import BackendApplicationClient
from requests_oauthlib import OAuth2Session
from dotenv import load_dotenv
import json
import postgrest.exceptions
import uuid
from werkzeug.utils import secure_filename
from enum import Enum

load_dotenv()

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "http://localhost:3000"}})

# Set up logging
logging.basicConfig(level=logging.DEBUG)

# Initialize Supabase client
url: str = os.environ.get("SUPABASE_URL")
key: str = os.environ.get("SUPABASE_KEY")
supabase: Client = create_client(url, key)

# Nettskjema API configuration
NETTSKJEMA_API_URL = "https://api.nettskjema.no/v3"
NETTSKJEMA_AUTH_URL = "https://authorization.nettskjema.no/oauth2/token"
NETTSKJEMA_CLIENT_ID = os.environ.get("NETTSKJEMA_CLIENT_ID")
NETTSKJEMA_CLIENT_SECRET = os.environ.get("NETTSKJEMA_CLIENT_SECRET")
NETTSKJEMA_FORM_ID = os.environ.get("NETTSKJEMA_FORM_ID")

class CompletionStatus(Enum):
    IN_PROGRESS = 'In Progress'
    COMPLETED_APPROVED = 'Completed and Approved'
    STOPPED = 'Stopped'
    NOT_APPROVED = 'Not Approved'

class DateTimeEncoder(json.JSONEncoder):
    def default(self, o):
        if isinstance(o, datetime):
            return o.isoformat()
        return super().default(o)

def get_oauth_session():
    client = BackendApplicationClient(client_id=NETTSKJEMA_CLIENT_ID)
    oauth = OAuth2Session(client=client)
    token = oauth.fetch_token(
        token_url=NETTSKJEMA_AUTH_URL,
        client_id=NETTSKJEMA_CLIENT_ID,
        client_secret=NETTSKJEMA_CLIENT_SECRET
    )
    return oauth

def get_nettskjema_data(form_id):
    oauth = get_oauth_session()
    response = oauth.get(f"{NETTSKJEMA_API_URL}/form/{form_id}/answers")
    response.raise_for_status()
    
    submissions = {}
    for line in response.iter_lines():
        if line:
            try:
                answer = json.loads(line.decode('utf-8'))
                submission_id = answer['submissionId']
                if submission_id not in submissions:
                    submissions[submission_id] = []
                submissions[submission_id].append(answer)
            except json.JSONDecodeError as e:
                logging.error(f"Error decoding JSON line: {e}")
                logging.error(f"Problematic line: {line}")
    
    return list(submissions.values())

def get_form_elements(form_id):
    oauth = get_oauth_session()
    response = oauth.get(f"{NETTSKJEMA_API_URL}/form/{form_id}/elements")
    response.raise_for_status()
    elements = response.json()
    
    # Create a mapping of element texts to their IDs
    element_mapping = {element['text']: element['elementId'] for element in elements}
    logging.info(f"Element mapping: {json.dumps(element_mapping, indent=2)}")
    
    return elements

def transform_submission_to_task(submission, element_mapping):
    task = {
        "submission_id": submission[0]['submissionId'],
        "title": "",
        "owner": "",
        "description": "",
        "relevance_for_bi": "",
        "need_for_course": "",
        "target_group": "",
        "growth_potential": "",
        "faculty_resources": "",
        "stage": "Idea Description",
        "completion_status": CompletionStatus.IN_PROGRESS.value,
        "attachment_url": None
    }
    
    for answer in submission:
        element_id = str(answer.get('elementId'))
        if element_id == str(element_mapping.get('Idea title')):
            task['title'] = answer.get('textAnswer', '')
        elif element_id == str(element_mapping.get('Idea Owner')):
            task['owner'] = answer.get('textAnswer', '')
        elif element_id == str(element_mapping.get('Briefly describe the idea:')):
            task['description'] = answer.get('textAnswer', '')
        elif element_id == str(element_mapping.get('Why is this relevant for BI?')):
            task['relevance_for_bi'] = answer.get('textAnswer', '')
        elif element_id == str(element_mapping.get('Why does individuals and/or organizations need such a course/idea?')):
            task['need_for_course'] = answer.get('textAnswer', '')
        elif element_id == str(element_mapping.get('What would be the relevant target group?')):
            task['target_group'] = answer.get('textAnswer', '')
        elif element_id == str(element_mapping.get('What are your thoughts on the future growth potential of the market for this course/idea?')):
            task['growth_potential'] = answer.get('textAnswer', '')
        elif element_id == str(element_mapping.get('Faculty resources – which academic departments should be involved?')):
            task['faculty_resources'] = answer.get('textAnswer', '')
        elif answer.get('answerAttachmentId'):
            task['attachment_url'] = f"https://api.nettskjema.no/v3/attachment/{answer.get('answerAttachmentId')}"
    
    return task

@app.route('/tasks', methods=['GET'])
def get_tasks():
    response = supabase.table('tasks').select('*').execute()
    tasks = response.data
    return jsonify(tasks)

@app.route('/tasks', methods=['POST'])
def add_task():
    new_task = request.json
    if 'id' in new_task:
        del new_task['id']  # Remove the id if it's present
    new_task['completion_status'] = CompletionStatus.IN_PROGRESS.value
    response = supabase.table('tasks').insert(new_task).execute()
    return jsonify(response.data[0]), 201

@app.route('/tasks/<int:task_id>', methods=['PUT'])
def update_task(task_id):
    updated_data = request.json
    try:
        # Get the current task data
        current_task = supabase.table('tasks').select('*').eq('id', task_id).execute()
        if not current_task.data:
            return jsonify({'error': 'Task not found'}), 404
        
        # Merge the updated data with the current task data
        task_data = current_task.data[0]
        task_data.update(updated_data)
        
        # If the stage is being updated to 'Completed', update completion_status
        if 'stage' in updated_data and updated_data['stage'] == 'Completed':
            task_data['completion_status'] = CompletionStatus.COMPLETED_APPROVED.value
        
        # Update the task
        response = supabase.table('tasks').update(task_data).eq('id', task_id).execute()
        if response.data:
            return jsonify(response.data[0]), 200
        else:
            return jsonify({'error': 'Failed to update task'}), 500
    except Exception as e:
        print(f"Error updating task: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/tasks/<int:task_id>', methods=['DELETE'])
def delete_task(task_id):
    supabase.table('tasks').delete().eq('id', task_id).execute()
    return '', 204

@app.route('/meetings', methods=['GET'])
def get_meetings():
    response = supabase.table('meetings').select('*').execute()
    meetings = response.data
    
    for meeting in meetings:
        # Fetch tasks for this meeting, including the stage_at_meeting
        tasks_response = supabase.table('meeting_tasks')\
            .select('task_id, stage_at_meeting, tasks(*)')\
            .eq('meeting_id', meeting['id'])\
            .execute()
        
        meeting['tasks'] = []
        for item in tasks_response.data:
            task = item['tasks']
            task['stage_at_meeting'] = item['stage_at_meeting']
            meeting['tasks'].append(task)
    
    return jsonify(meetings), 200

@app.route('/meetings', methods=['POST'])
def add_meeting():
    new_meeting = request.json
    print("Received meeting data:", new_meeting)

    meeting_data = {
        'number': new_meeting.get('number'),
        'date': new_meeting.get('date'),
        'location': new_meeting.get('location')
    }

    if not meeting_data['number'] or not meeting_data['date']:
        return jsonify({'error': 'Meeting number and date are required'}), 400

    if isinstance(meeting_data['date'], str):
        try:
            # Parse the date string to a datetime object
            date_obj = datetime.fromisoformat(meeting_data['date'].replace('Z', '+00:00'))
            # Convert the datetime object back to an ISO 8601 string
            meeting_data['date'] = date_obj.isoformat()
        except ValueError:
            return jsonify({'error': 'Invalid date format. Please use ISO format (YYYY-MM-DDTHH:MM:SS.sssZ)'}), 400

    try:
        print("Inserting meeting data:", meeting_data)
        response = supabase.table('meetings').insert(meeting_data).execute()
        print("Supabase response:", response)
        if response.data:
            new_meeting = response.data[0]
            return_data = {
                'id': new_meeting['id'],
                'number': new_meeting['number'],
                'date': new_meeting['date'],
                'location': new_meeting['location']
            }
            return jsonify(return_data), 201
        else:
            return jsonify({'error': 'Failed to create meeting'}), 500
    except Exception as e:
        print("Error:", str(e))
        return jsonify({'error': str(e)}), 400
    
@app.route('/meetings/<int:meeting_id>', methods=['PUT'])
def update_meeting(meeting_id):
    updated_data = request.json
    response = supabase.table('meetings').update(updated_data).eq('id', meeting_id).execute()
    return jsonify(response.data[0])

@app.route('/meetings/<int:meeting_id>', methods=['DELETE'])
def delete_meeting(meeting_id):
    supabase.table('meetings').delete().eq('id', meeting_id).execute()
    return '', 204

@app.route('/meetings/<int:meeting_id>/tasks', methods=['POST'])
def add_task_to_meeting(meeting_id):
    data = request.json
    task_id = data['task_id']
    
    # Fetch the meeting
    meeting_response = supabase.table('meetings').select('*').eq('id', meeting_id).execute()
    if not meeting_response.data:
        return jsonify({"error": "Meeting not found"}), 404
    meeting = meeting_response.data[0]
    
    # Fetch the task
    task_response = supabase.table('tasks').select('*').eq('id', task_id).execute()
    if not task_response.data:
        return jsonify({"error": "Task not found"}), 404
    task = task_response.data[0]
    
    if not meeting.get('is_completed', False):
        # Get the current number of tasks in the meeting
        current_tasks_response = supabase.table('meeting_tasks').select('task_id').eq('meeting_id', meeting_id).execute()
        new_order = len(current_tasks_response.data) + 1
        
        # Add the task to the meeting, including the current stage
        meeting_task_data = {
            'meeting_id': meeting_id,
            'task_id': task_id,
            'task_order': new_order,
            'stage_at_meeting': task['stage']  # Store the current stage of the task
        }
        supabase.table('meeting_tasks').insert(meeting_task_data).execute()
        
        return jsonify({"message": "Task added to meeting", "stage_at_meeting": task['stage']}), 201
    else:
        return jsonify({"error": "Cannot add tasks to a completed meeting"}), 400
    
@app.route('/meetings/<int:meeting_id>/complete', methods=['PUT'])
def complete_meeting(meeting_id):
    response = supabase.table('meetings').update({'is_completed': True}).eq('id', meeting_id).execute()
    if response.data:
        return jsonify({"message": "Meeting marked as completed"}), 200
    else:
        return jsonify({"error": "Meeting not found"}), 404

@app.route('/meetings/<int:meeting_id>/reorder', methods=['PUT'])
def reorder_meeting_tasks(meeting_id):
    data = request.json
    new_order = data['new_order']  # List of task IDs in the new order
    
    meeting_response = supabase.table('meetings').select('*').eq('id', meeting_id).execute()
    if not meeting_response.data:
        return jsonify({"error": "Meeting not found"}), 404
    meeting = meeting_response.data[0]
    
    if not meeting.get('is_completed', False):
        for index, task_id in enumerate(new_order, start=1):
            # Update the order in meeting_tasks
            supabase.table('meeting_tasks').update({'order': index}).eq('meeting_id', meeting_id).eq('task_id', task_id).execute()
        
        return jsonify({"message": "Tasks reordered successfully"}), 200
    else:
        return jsonify({"error": "Cannot reorder tasks in a completed meeting"}), 400

@app.route('/meetings/<int:meeting_id>/tasks/<int:task_id>', methods=['DELETE'])
def remove_task_from_meeting(meeting_id, task_id):
    supabase.table('meeting_tasks').delete().eq('meeting_id', meeting_id).eq('task_id', task_id).execute()
    return '', 204

@app.route('/meetings/<int:meeting_id>/tasks/<int:task_id>', methods=['PUT'])
def update_meeting_task(meeting_id, task_id):
    data = request.json
    if 'minutes' in data:
        response = supabase.table('meeting_tasks').update({'minutes': data['minutes']}).eq('meeting_id', meeting_id).eq('task_id', task_id).execute()
    return jsonify(response.data[0]), 200

@app.route('/meetings/<int:meeting_id>/generate_report', methods=['GET'])
def generate_report(meeting_id):
    meeting_response = supabase.table('meetings').select('*').eq('id', meeting_id).execute()
    if not meeting_response.data:
        return jsonify({'error': 'Meeting not found'}), 404
    
    meeting = meeting_response.data[0]
    
    # Update this query to include stage_at_meeting
    tasks_response = supabase.table('meeting_tasks').select('tasks(*), stage_at_meeting').eq('meeting_id', meeting_id).execute()
    tasks = tasks_response.data

    document = Document()
    
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    title = document.add_paragraph("Innovation Board Executive")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(14)

    meeting_date = meeting['date']
    details = [
        ("Place", "A4Y-117"),
        ("Date and time", meeting_date)
    ]
    for key, value in details:
        p = document.add_paragraph()
        p.add_run(f"{key}\n").bold = True
        p.add_run(value)

    members = [
        ("Lars Olsen", "Dean Executive"),
        ("Cecilie Asting", "Associate Dean Bachelor of Management"),
        ("Geir Høidal Bjønnes", "Associate Dean Master of Management"),
        ("Lise Hammergren", "Executive Vice President, BI Executive"),
        ("Gry Varre", "Manager Open Enrolment"),
        ("Wenche Martinussen", "Director Sales/Marketing"),
        ("Tonje Omland", "Manager Programme administration - Executive"),
        ("Jørgen Bjørnson Aanderaa", "Director Learning Center"),
        ("Line Lervik-Olsen", "Head of Department, Marketing"),
        ("Eivind Furseth", "Head of Department, Accounting and Operations Management")
    ]
    
    document.add_paragraph("Members", style='Heading 1')
    for name, title in members:
        document.add_paragraph(f"{name} {title}")

    document.add_paragraph("Programme Administration / Secretariat", style='Heading 1')
    secretariat = [
        ("Nora Iversen Røed", "Adviser"),
        ("Haakon Tveter", "Senior Adviser")
    ]
    for name, title in secretariat:
        document.add_paragraph(f"{name} {title}")

    document.add_page_break()

    document.add_paragraph("Agenda", style='Heading 1')
    for index, task in enumerate(tasks, start=1):
        document.add_paragraph(f"{task['tasks']['casenumber']} {task['tasks']['title']} page {index}")

    for task in tasks:
        document.add_page_break()

        document.add_paragraph("Proposal – New course idea", style='Heading 2')
        
        headings = [
            ("Item number", task['tasks']['casenumber']),
            ("Idea title", task['tasks']['title']),
            ("Idea owner", task['tasks']['owner']),
            ("Stage at meeting", task['stage_at_meeting']),
            ("Current stage", task['tasks']['stage']),
            ("Completion status", task['tasks']['completion_status']),
            ("Briefly describe the idea", task['tasks']['description']),
            ("Why is this relevant for BI?", task['tasks'].get('relevance_for_bi', '')),
            ("Why does individuals and/or organizations need such a course/idea?", task['tasks'].get('need_for_course', '')),
            ("What would be the relevant target group?", task['tasks'].get('target_group', '')),
            ("What are your thoughts on the future growth potential of the market for this course/idea?", task['tasks'].get('growth_potential', '')),
            ("Faculty resources – which academic departments should be involved?", task['tasks'].get('faculty_resources', ''))
        ]

        for heading, content in headings:
            document.add_paragraph(heading, style='Heading 3')
            document.add_paragraph(content)

    f = BytesIO()
    document.save(f)
    f.seek(0)

    return send_file(f, as_attachment=True, download_name='innovation_board_sakspapirer.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/tasks/<int:task_id>/upload', methods=['POST'])
def upload_file(task_id):
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file:
        original_filename = file.filename
        filename = secure_filename(original_filename)
        file_extension = os.path.splitext(filename)[1]
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        
        try:
            # Upload to Supabase Storage
            file_content = file.read()
            upload_response = supabase.storage.from_('task-attachments').upload(unique_filename, file_content)
            
            # Get public URL
            public_url = supabase.storage.from_('task-attachments').get_public_url(unique_filename)
            
            # Create attachment object
            new_attachment = {
                'url': public_url,
                'filename': original_filename
            }
            
            # Fetch current task data
            task_response = supabase.table('tasks').select('attachments').eq('id', task_id).execute()
            if task_response.data:
                current_attachments = task_response.data[0].get('attachments', []) or []
            else:
                current_attachments = []
            
            # Append new attachment
            updated_attachments = current_attachments + [new_attachment]
            
            # Update task with new attachments
            update_response = supabase.table('tasks').update({'attachments': updated_attachments}).eq('id', task_id).execute()
            
            return jsonify({
                'message': 'File uploaded successfully',
                'attachment': new_attachment
            }), 200
        except Exception as e:
            logging.error(f"Error uploading file: {str(e)}")
            return jsonify({'error': f'Error uploading file: {str(e)}'}), 500

@app.route('/meetings/<int:meeting_id>/generate_minutes', methods=['GET'])
def generate_minutes(meeting_id):
    meeting_response = supabase.table('meetings').select('*').eq('id', meeting_id).execute()
    if not meeting_response.data:
        return jsonify({'error': 'Meeting not found'}), 404
    
    meeting = meeting_response.data[0]
    
    tasks_response = supabase.table('meeting_tasks').select('tasks(*), minutes, stage_at_meeting').eq('meeting_id', meeting_id).execute()
    tasks = tasks_response.data

    document = Document()
    
    title = document.add_paragraph(f"Møtereferat for møte {meeting['number']}")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    date_paragraph = document.add_paragraph(f"Dato: {meeting['date']}")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for task in tasks:
        item_title = document.add_paragraph(f"{task['tasks']['casenumber']} - {task['tasks']['title']}")
        item_title.runs[0].bold = True
        
        document.add_paragraph(f"Stage at meeting: {task['stage_at_meeting']}")
        document.add_paragraph(f"Current stage: {task['tasks']['stage']}")
        document.add_paragraph(f"Completion status: {task['tasks']['completion_status']}")
        
        if task['minutes']:
            document.add_paragraph(task['minutes'])
        else:
            document.add_paragraph("Ingen referat tilgjengelig for dette punktet.")
        
        document.add_paragraph()

    f = BytesIO()
    document.save(f)
    f.seek(0)

    return send_file(f, as_attachment=True, download_name=f'meeting_{meeting_id}_minutes.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/import-nettskjema', methods=['POST'])
def import_nettskjema():
    imported_tasks = []
    try:
        logging.info("Starting Nettskjema import process")
        
        # Get form elements
        form_elements = get_form_elements(NETTSKJEMA_FORM_ID)
        logging.info(f"Form elements: {json.dumps(form_elements, indent=2)}")
        
        submissions = get_nettskjema_data(NETTSKJEMA_FORM_ID)
        logging.info(f"Received {len(submissions)} submissions from Nettskjema")
        
        if not submissions:
            logging.warning("No submissions received from Nettskjema")
            return jsonify({"message": "No submissions to import"}), 200

        existing_submission_ids = set(task['submission_id'] for task in supabase.table('tasks').select('submission_id').execute().data)

        form_elements = get_form_elements(NETTSKJEMA_FORM_ID)
        element_mapping = {element['text']: element['elementId'] for element in form_elements}
        logging.info(f"Element mapping: {json.dumps(element_mapping, indent=2)}")

        for submission in submissions:
            logging.info(f"Processing submission: {submission[0]['submissionId']}")
            task = transform_submission_to_task(submission, element_mapping)
            logging.info(f"Transformed task: {json.dumps(task, indent=2)}")

            # Log each field separately
            for key, value in task.items():
                logging.info(f"{key}: {value}")

            if task['submission_id'] in existing_submission_ids:
                logging.info(f"Skipping duplicate submission: {task['submission_id']}")
                continue

            # Handle attachment upload
            if task['attachment_url']:
                try:
                    response = requests.get(task['attachment_url'])
                    response.raise_for_status()
                    file_content = BytesIO(response.content)
                    file_name = f"{uuid.uuid4()}{os.path.splitext(response.headers.get('Content-Disposition', ''))[1]}"
                    
                    # Upload to Supabase Storage
                    upload_response = supabase.storage.from_('task-attachments').upload(file_name, file_content)
                    
                    # Get public URL
                    public_url = supabase.storage.from_('task-attachments').get_public_url(file_name)
                    
                    task['attachment_url'] = public_url
                    logging.info(f"Uploaded attachment: {public_url}")
                except Exception as e:
                    logging.error(f"Error uploading attachment: {str(e)}")
                    task['attachment_url'] = None

            try:
                response = supabase.table('tasks').insert(task).execute()
                if response.data:
                    imported_tasks.extend(response.data)
                    logging.info(f"Successfully imported task: {task['title']}")
                    logging.info(f"Imported task data: {json.dumps(response.data[0], indent=2)}")
                else:
                    logging.warning(f"No data returned when inserting task: {task}")
            except Exception as e:
                logging.error(f"Error inserting task into database: {str(e)}", exc_info=True)

        logging.info(f"Successfully imported {len(imported_tasks)} tasks")
        return jsonify({
            "message": f"Successfully imported {len(imported_tasks)} tasks",
            "imported_tasks": imported_tasks
        }), 200
    except requests.RequestException as e:
        logging.error(f"Error fetching data from Nettskjema: {str(e)}", exc_info=True)
        return jsonify({"error": f"Error fetching data from Nettskjema: {str(e)}"}), 500
    except Exception as e:
        logging.error(f"Error processing import: {str(e)}", exc_info=True)
        return jsonify({"error": f"Error processing import: {str(e)}"}), 500

@app.route('/tasks/<int:task_id>/meetings', methods=['GET'])
def get_task_meeting_history(task_id):
    try:
        # Query to get all meetings for this task, along with the stage at each meeting
        response = supabase.table('meeting_tasks')\
            .select('meetings(id, date, number), stage_at_meeting')\
            .eq('task_id', task_id)\
            .execute()
        
        if not response.data:
            return jsonify([]), 200
        
        # Format the response data
        meeting_history = []
        for entry in response.data:
            meeting_info = entry['meetings']
            meeting_history.append({
                'date': meeting_info['date'],
                'number': meeting_info['number'],
                'stage_at_meeting': entry['stage_at_meeting']
            })
        
        # Sort the meeting history by date
        meeting_history.sort(key=lambda x: x['date'])
        
        return jsonify(meeting_history), 200
    except Exception as e:
        print(f"Error fetching meeting history: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/tasks/<int:task_id>/status', methods=['PUT'])
def update_task_status(task_id):
    new_status = request.json.get('status')
    if new_status not in [status.value for status in CompletionStatus]:
        return jsonify({'error': 'Invalid status'}), 400
    
    try:
        response = supabase.table('tasks').update({'completion_status': new_status}).eq('id', task_id).execute()
        if response.data:
            return jsonify(response.data[0]), 200
        else:
            return jsonify({'error': 'Failed to update task status'}), 500
    except Exception as e:
        print(f"Error updating task status: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.errorhandler(404)
def not_found(error):
    return jsonify({"error": "Not found"}), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({"error": "Internal server error"}), 500

if __name__ == '__main__':
    app.run(debug=True)